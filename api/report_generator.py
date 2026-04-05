"""
Генерация отчётов в форматах PDF, DOCX, JSON.
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from core.models import CalcResponse, PointResult
from core.constants import OCTAVE_FREQS


def generate_json_report(response: CalcResponse, project_name: str = "Акустик") -> str:
    """Отчёт в формате JSON."""
    report = {
        "project": project_name,
        "algorithm": response.metadata.get("algorithm", ""),
        "summary": {
            "sources_count": response.metadata.get("sources_count", 0),
            "points_count": response.metadata.get("points_count", 0),
            "exceeded_count": sum(1 for r in response.results if r.exceeded),
        },
        "results": [r.model_dump() for r in response.results],
    }
    return json.dumps(report, ensure_ascii=False, indent=2)


def generate_docx_report(response: CalcResponse,
                          output_path: str,
                          project_name: str = "Акустик") -> str:
    """Отчёт в формате DOCX (python-docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return ""

    doc = Document()

    # Заголовок
    title = doc.add_heading(f"Акустический расчёт: {project_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Алгоритм: {response.metadata.get('algorithm', 'СП 51.13330.2011')}"
    )
    doc.add_paragraph(
        f"Источников шума: {response.metadata.get('sources_count', '—')}, "
        f"расчётных точек: {response.metadata.get('points_count', '—')}"
    )

    doc.add_heading("Результаты расчёта", 1)

    # Таблица результатов
    headers = ["РТ", "Период"] + [str(f) for f in OCTAVE_FREQS] + ["Lэкв,дБА", "ПДУ,дБА", "ΔL,дБА", "Оценка"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_row = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_row[i].text = h

    for result in response.results:
        row = table.add_row().cells
        row[0].text = result.point_id
        row[1].text = "День" if result.period == "day" else "Ночь"
        for i, l in enumerate(result.l_octave):
            row[2 + i].text = f"{l:.1f}"
        row[10].text = f"{result.l_a_eq:.1f}"
        row[11].text = f"{result.pdu_eq:.0f}"
        exc = result.exceedance_eq
        row[12].text = f"{'+' if exc > 0 else ''}{exc:.1f}"
        row[13].text = "ПРЕВЫШЕНИЕ!" if result.exceeded else "Норма"

    doc.save(output_path)
    return output_path


def generate_pdf_report(response: CalcResponse,
                         output_path: str,
                         project_name: str = "Акустик") -> str:
    """Отчёт в формате PDF (reportlab)."""
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return ""

    styles = getSampleStyleSheet()

    doc = SimpleDocTemplate(
        output_path,
        pagesize=landscape(A4),
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
    )
    story = []

    # Заголовок
    story.append(Paragraph(f"Акустический расчет: {project_name}", styles["Title"]))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        f"Алгоритм: {response.metadata.get('algorithm', 'СП 51.13330.2011')} | "
        f"Источников: {response.metadata.get('sources_count', '—')} | "
        f"Точек: {response.metadata.get('points_count', '—')}",
        styles["Normal"]
    ))
    story.append(Spacer(1, 0.5 * cm))

    # Таблица
    freq_headers = [str(f) for f in OCTAVE_FREQS]
    headers = ["РТ", "Период"] + freq_headers + ["Lэкв", "ПДУ", "Delta", "Оценка"]
    data = [headers]

    for result in response.results:
        exc = result.exceedance_eq
        row = (
            [result.point_id, "День" if result.period == "day" else "Ночь"]
            + [f"{l:.1f}" for l in result.l_octave]
            + [
                f"{result.l_a_eq:.1f}",
                f"{result.pdu_eq:.0f}",
                f"{'+' if exc > 0 else ''}{exc:.1f}",
                "ПРЕВЫШЕНИЕ" if result.exceeded else "Норма",
            ]
        )
        data.append(row)

    col_widths = [1.5 * cm, 1.2 * cm] + [1.5 * cm] * 8 + [1.5 * cm, 1.2 * cm, 1.5 * cm, 2.0 * cm]
    t = Table(data, colWidths=col_widths, repeatRows=1)

    style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.black),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ])

    # Подсветить превышения красным
    for row_idx, result in enumerate(response.results, start=1):
        if result.exceeded:
            style.add("BACKGROUND", (-1, row_idx), (-1, row_idx), colors.salmon)
            style.add("TEXTCOLOR", (-1, row_idx), (-1, row_idx), colors.darkred)

    t.setStyle(style)
    story.append(t)

    doc.build(story)
    return output_path
